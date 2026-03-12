#!/usr/bin/env python3.11
"""
Generate Boomi-branded Customer Use-Case Submission Template (.docx)
Single continuous table — no floating boxes.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Boomi Brand Colors (verified from official Boomi SVG assets) ──────────────
# #003D58 — primary navy (logo icon bg, boomi-1.svg)
# #072B55 — dark navy (wordmark text, brand.svg)
# #FF7C66 — coral accent (both official SVGs)
# #00A9CE — secondary teal (Boomi marketing standard)
BOOMI_NAVY   = RGBColor(0x00, 0x3D, 0x58)   # primary navy — confirmed
BOOMI_DARK   = RGBColor(0x07, 0x2B, 0x55)   # dark navy wordmark — confirmed
BOOMI_TEAL   = RGBColor(0x00, 0xA9, 0xCE)   # secondary teal
BOOMI_CORAL  = RGBColor(0xFF, 0x7C, 0x66)   # coral accent — confirmed
BOOMI_MID    = RGBColor(0x00, 0x55, 0x80)   # mid-tone blue (derived)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_DARK    = RGBColor(0x07, 0x2B, 0x55)   # use confirmed dark navy for body text
TEXT_MID     = RGBColor(0x47, 0x55, 0x69)
CRITICAL_RED = RGBColor(0xDC, 0x26, 0x26)
NICE_GREEN   = RGBColor(0x16, 0xA3, 0x4A)

OUTPUT_PATH = ('/mnt/c/users/BrianMerrick/Documents/Dev/ClaudeCode/'
               'boomicompanion_template_workspace/business-demo/boomi/'
               'boomi-usecase-submission-template.docx')

COL_LABEL = Inches(1.8)
COL_VALUE = Inches(5.5)


# ─── Low-level helpers ─────────────────────────────────────────────────────────

def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    # remove existing shd
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(shd)


def set_table_borders(table, color='CBD5E1', sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.append(tblBorders)


def tight(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    pPr.append(spacing)


def run(para, text, bold=False, italic=False, color=None, size=9):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Calibri'
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return r


def merge_row_cells(row):
    """Merge all cells in a row into one."""
    row.cells[0].merge(row.cells[1])


def fix_widths(table):
    for row in table.rows:
        if len(row.cells) == 2:
            row.cells[0].width = COL_LABEL
            row.cells[1].width = COL_VALUE


# ─── Row builders ─────────────────────────────────────────────────────────────

def add_section_header(table, title, subtitle='', bg='003D58', page_break=False):
    """Full-width merged header row."""
    row = table.add_row()
    merge_row_cells(row)
    cell = row.cells[0]
    cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    tight(p)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    if page_break:
        p.paragraph_format.page_break_before = True
    run(p, f'  {title}', bold=True, color=WHITE, size=11)
    if subtitle:
        run(p, f'   {subtitle}', bold=False, color=BOOMI_TEAL, size=8)


def add_field(table, label, placeholder):
    row = table.add_row()
    lc, vc = row.cells[0], row.cells[1]
    cell_bg(lc, 'E6F4F9')
    cell_bg(vc, 'FFFFFF')
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    lp = lc.paragraphs[0]
    tight(lp)
    run(lp, label, bold=True, color=BOOMI_NAVY, size=8.5)

    vp = vc.paragraphs[0]
    tight(vp)
    run(vp, placeholder, italic=True, color=TEXT_MID, size=8.5)


def add_multifield(table, label, lines):
    row = table.add_row()
    lc, vc = row.cells[0], row.cells[1]
    cell_bg(lc, 'E6F4F9')
    cell_bg(vc, 'FFFFFF')
    lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    vc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    lp = lc.paragraphs[0]
    tight(lp)
    run(lp, label, bold=True, color=BOOMI_NAVY, size=8.5)

    for i, line in enumerate(lines):
        vp = vc.paragraphs[0] if i == 0 else vc.add_paragraph()
        tight(vp)
        run(vp, line, italic=True, color=TEXT_MID, size=8.5)


def add_priority(table):
    row = table.add_row()
    lc, vc = row.cells[0], row.cells[1]
    cell_bg(lc, 'E6F4F9')
    cell_bg(vc, 'FFFFFF')
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    lp = lc.paragraphs[0]
    tight(lp)
    run(lp, 'Priority', bold=True, color=BOOMI_NAVY, size=8.5)

    vp = vc.paragraphs[0]
    tight(vp)
    run(vp, '☐  Critical    ', bold=True, color=CRITICAL_RED, size=8.5)
    run(vp, '☐  Nice-to-have', bold=True, color=NICE_GREEN, size=8.5)


def add_use_case_fields(table, number):
    add_section_header(table, f'USE CASE {number}', page_break=(number > 1))
    add_field(table,      'Title',              '[Short descriptive title]')
    add_field(table,      'Objective',          '[Problem this solves / why it matters]')
    add_field(table,      'Contact',            '[Name, Role, Email]')
    add_priority(table)
    add_field(table,      'Trigger',            '[Event / schedule / API call / file drop]')
    add_multifield(table, 'Source → Target',    ['Source(s): [System A]', 'Target(s): [System B]'])
    add_field(table,      'Data Elements',      '[customerId, orderId, amount, status …]')
    add_field(table,      'Volume & Frequency', '[10k/day · peak 200/min · hourly batch]')
    add_multifield(table, 'KPIs',               ['KPI 1: [metric + target]', 'KPI 2: [metric + target]', 'KPI 3: [metric + target]'])
    add_field(table,      'Success Criteria',   '[Pass/fail with numbers]')
    add_multifield(table, 'Happy Path',         ['1. [Trigger]', '2. [Boomi receives / polls]', '3. [Transform / enrich]', '4. [Write to target]', '5. [Target confirms → ack]'])
    add_multifield(table, 'Failure Modes',      ['[Source 500]  →  [retry 3×, alert]', '[Validation fail]  →  [error queue, notify]', '[Target down]  →  [queue, auto-retry, escalate]'])
    add_field(table,      'Data Flow',          '[Source → Boomi → Target  |  diagram link]')
    add_field(table,      'Security',           '[PII? Encryption? SOC2 / HIPAA / masking?]')
    add_field(table,      'Dependencies',       '[Credentials, firewall rules, certs, approvals]')
    add_multifield(table, 'Acceptance Tests',   ['Test 1: [Happy path end-to-end]', 'Test 2: [Malformed record → error queue]', 'Test 3: [Target failure → alert fired]'])
    add_field(table,      'Business Impact',    '[Saves X hrs/month · reduces cycle time by Y%]')
    add_field(table,      'Notes / Attachments','[Sample payloads, schemas, screenshots, SLAs]')


# ─── Main ──────────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)
    section.top_margin    = Inches(0.65)
    section.bottom_margin = Inches(0.65)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9)

    # ── Page header with Boomi logo ────────────────────────────────────────────
    section.header_distance = Inches(0.25)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(hp)
    logo_path = ('/mnt/c/users/BrianMerrick/Documents/Dev/ClaudeCode/'
                 'boomicompanion_template_workspace/business-demo/boomi/boomi-logo-official.png')
    hr = hp.add_run()
    hr.add_picture(logo_path, height=Inches(0.35))

    # ── Title paragraph ────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(2)
    run(tp, 'Customer Use-Case Submission Template', bold=True, color=BOOMI_NAVY, size=16)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after  = Pt(6)
    run(sp, 'Up to 3 use-cases', color=BOOMI_MID, size=9)

    # ── Instructions ──────────────────────────────────────────────────────────
    inst_label = doc.add_paragraph()
    inst_label.paragraph_format.space_before = Pt(0)
    inst_label.paragraph_format.space_after  = Pt(1)
    run(inst_label, 'Instructions', bold=True, color=BOOMI_NAVY, size=10)

    instructions = [
        ('1.', 'Complete up to 3 use-case blocks (one per integration / flow).'),
        ('2.', 'Mark Priority as Critical or Nice-to-have.'),
        ('3.', 'KPIs should be specific — e.g., "< 2% error rate", "< 5 min latency", "6 steps → 1".'),
        ('4.', 'Happy Path = ideal flow. Failure Modes = errors and desired behavior.'),
        ('5.', 'Attach sample payloads, schemas, or screenshots separately if helpful.'),
    ]
    for num, text in instructions:
        ip = doc.add_paragraph()
        ip.paragraph_format.space_before  = Pt(0)
        ip.paragraph_format.space_after   = Pt(0)
        ip.paragraph_format.left_indent   = Inches(0.1)
        run(ip, f'{num}  ', bold=True, color=BOOMI_TEAL, size=8.5)
        run(ip, text, color=TEXT_DARK, size=8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Single master table ────────────────────────────────────────────────────
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    set_table_borders(table, color='CBD5E1', sz=4)

    # tblGrid
    tbl = table._tbl
    tblGrid = OxmlElement('w:tblGrid')
    for w in [1800, 5500]:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    for old in tbl.findall(qn('w:tblGrid')):
        tbl.remove(old)
    tbl.insert(1, tblGrid)

    # 3 use case blocks
    for i in range(1, 4):
        add_use_case_fields(table, i)

    # Example block — mirrors every field in add_use_case_fields
    add_section_header(table, 'EXAMPLE — Filled Snippet', bg='072B55', page_break=True)

    def ex_field(label, value):
        row = table.add_row()
        lc, vc = row.cells[0], row.cells[1]
        cell_bg(lc, 'E6F4F9')
        cell_bg(vc, 'FFFFFF')
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lc.paragraphs[0]; tight(lp)
        run(lp, label, bold=True, color=BOOMI_MID, size=8.5)
        vp = vc.paragraphs[0]; tight(vp)
        is_crit = value.startswith('☑')
        run(vp, value, bold=is_crit, color=CRITICAL_RED if is_crit else TEXT_DARK, size=8.5)

    def ex_multifield(label, lines):
        row = table.add_row()
        lc, vc = row.cells[0], row.cells[1]
        cell_bg(lc, 'E6F4F9')
        cell_bg(vc, 'FFFFFF')
        lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        vc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        lp = lc.paragraphs[0]; tight(lp)
        run(lp, label, bold=True, color=BOOMI_MID, size=8.5)
        for i, line in enumerate(lines):
            vp = vc.paragraphs[0] if i == 0 else vc.add_paragraph()
            tight(vp)
            run(vp, line, color=TEXT_DARK, size=8.5)

    ex_field('Title',            'Sync new e-commerce orders to ERP')
    ex_field('Objective',        'Push orders placed on storefront to ERP immediately to start fulfillment.')
    ex_field('Contact',          'Jane Smith, VP Operations, jane.smith@acme.com')
    ex_field('Priority',         '☑  Critical')
    ex_field('Trigger',          'Order.created webhook event from Shopify storefront')
    ex_multifield('Source → Target', [
        'Source(s): Shopify (webhook)',
        'Target(s): Oracle ERP (REST API)',
    ])
    ex_field('Data Elements',      'orderId, customerId, lineItems, amount, currency, shippingAddress, status')
    ex_field('Volume & Frequency', '~5k orders/day · peak 300/min during promotions · near-real-time')
    ex_multifield('KPIs', [
        'KPI 1: Order processing latency < 2 min (p95)',
        'KPI 2: Error rate < 1%',
        'KPI 3: Zero duplicate orders in ERP',
    ])
    ex_field('Success Criteria',   '95% of orders appear in ERP within 2 min; failed orders land in error queue within 30 s')
    ex_multifield('Happy Path', [
        '1. Shopify fires Order.created webhook to Boomi listener',
        '2. Boomi validates payload schema and deduplicates by orderId',
        '3. Enrich with customer pricebook from ERP reference data',
        '4. POST transformed order to Oracle ERP /orders endpoint',
        '5. ERP returns 201 Created → Boomi acknowledges webhook → success logged',
    ])
    ex_multifield('Failure Modes', [
        'ERP returns 503  →  retry with exponential backoff (3×), then persist to dead-letter queue and alert ops',
        'Schema validation fails  →  route to error queue, notify data-ops team via email',
        'Duplicate orderId detected  →  discard silently, log warning',
    ])
    ex_field('Data Flow',        'Shopify → Boomi Webhook Listener → Validate/Enrich → Oracle ERP API')
    ex_field('Security',         'PII present (name, address); TLS in transit; fields masked in logs; SOC 2 compliant')
    ex_field('Dependencies',     'Shopify webhook secret, Oracle ERP API key, firewall rule for Boomi IP range')
    ex_multifield('Acceptance Tests', [
        'Test 1: Place test order → verify it appears in ERP within 2 min with correct line items',
        'Test 2: Send malformed JSON payload → verify it lands in error queue, no ERP write',
        'Test 3: Take ERP offline → verify retry fires 3× then alert email received by ops',
    ])
    ex_field('Business Impact',  'Eliminates 40 hrs/month of manual order entry; reduces order-to-fulfillment cycle by 25%')
    ex_field('Notes / Attachments', 'Sample order payload: /docs/shopify-order-sample.json · ERP API spec: /docs/erp-api-v2.yaml')

    fix_widths(table)


    doc.save(OUTPUT_PATH)
    print(f'Saved: {OUTPUT_PATH}')


if __name__ == '__main__':
    build_document()
